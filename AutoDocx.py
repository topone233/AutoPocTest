import time
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

document = Document()
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Normal'].font.size = Pt(12)
document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

data = {
    '漏洞系统': 'Test Site',
    '官方网站': 'http://URL',
    '漏洞类型': 'SQL_Injection',
    '漏洞信息': 'information',
    '漏洞复现': 'review',
}


def writeup(title, body):
    # 添加标题
    head = document.add_heading('', 1)
    run = head.add_run(title)
    run.font.name = u'宋体'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 添加段落
    document.add_paragraph(body)


for i in data:
    title = i
    body = data[i]
    writeup(title, body)

time_list = int(time.time())
document.save(str(time_list) + '.docx')
print('Saved at ' + str(time_list) + '.docx')
